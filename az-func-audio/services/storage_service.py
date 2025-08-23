import os
import logging
from typing import Optional, Any
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from azure.core.exceptions import AzureError
from datetime import datetime, timedelta
from urllib.parse import urlparse

from config import AppConfig

logger = logging.getLogger(__name__)


class StorageServiceError(Exception):
    """Custom exception for storage service errors."""
    pass

class StorageService:
    def __init__(self, config: AppConfig, credential: Any = None, blob_service_client: BlobServiceClient = None) -> None:
        """Initialize the StorageService with config, optional credential, and blob service client."""
        self.config = config
        # Lazy import of DefaultAzureCredential to avoid import-time failures
        if credential is not None:
            self.credential = credential
        else:
            try:
                from azure.identity import DefaultAzureCredential
                self.credential = DefaultAzureCredential()
            except Exception:
                self.credential = None

        # Initialize blob service client
        self.blob_service_client = blob_service_client if blob_service_client is not None else BlobServiceClient(
            account_url=self.config.storage_account_url,
            credential=self.credential,
        )

    def upload_file(self, file_path: str, original_filename: str) -> str:
        """Upload a file to blob storage and return the blob URL."""
        try:
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )            # Generate blob name with date and nested structure including timestamp for uniqueness
            current_date = datetime.now().strftime("%Y-%m-%d")
            timestamp = datetime.now().strftime("%H%M%S_%f")[:-3]  # HHMMSS_milliseconds
            file_name_without_ext = os.path.splitext(original_filename)[0]
            # Include timestamp in both folder and filename to ensure uniqueness
            blob_name = f"{current_date}/{file_name_without_ext}_{timestamp}/{original_filename}"

            blob_client = container_client.get_blob_client(blob_name)

            # Upload the file
            logger.info(f"Uploading file to blob storage: {blob_name}")
            with open(file_path, "rb") as data:
                blob_client.upload_blob(data, overwrite=True)

            return blob_client.url

        except AzureError as e:
            logger.error(f"Azure storage error: {str(e)}")
            raise StorageServiceError(f"Azure storage error: {str(e)}") from e
        except Exception as e:
            logger.error(f"Error uploading file: {str(e)}")
            raise StorageServiceError(f"Error uploading file: {str(e)}") from e

    def upload_text(
        self, container_name: str, blob_name: str, text_content: str
    ) -> str:
        """Upload text content to blob storage and return the blob URL."""
        try:
            container_client = self.blob_service_client.get_container_client(
                container_name
            )
            blob_client = container_client.get_blob_client(blob_name)

            blob_client.upload_blob(text_content.encode("utf-8"), overwrite=True)
            return blob_client.url
        except Exception as e:
            logger.error(f"Error uploading text: {str(e)}")
            raise StorageServiceError(f"Error uploading text: {str(e)}") from e

    def generate_and_upload_pdf(self, analysis_text: str, blob_url: str) -> str:
        """Generate a PDF from analysis text and upload to blob storage. Return the blob URL."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            import io

            # Create PDF in memory
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)

            # Add content to PDF
            y = 750  # Starting y position
            for line in analysis_text.split("\n"):
                if y < 50:  # Start new page if near bottom
                    c.showPage()
                    y = 750
                c.drawString(50, y, line)
                y -= 15

            c.save()
            pdf_content = buffer.getvalue()

            # Upload PDF
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_url)

            blob_client.upload_blob(pdf_content, overwrite=True)
            return blob_client.url

        except Exception as e:
            logger.error(f"Error generating/uploading PDF: {str(e)}")
            raise StorageServiceError(f"Error generating/uploading PDF: {str(e)}") from e

    def generate_and_upload_docx(self, analysis_text: str, blob_url: str) -> str:
        """Generate a DOCX from analysis text and upload to blob storage. Return the blob URL."""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import io
            import re

            # Create DOCX in memory
            doc = Document()
            
            # Add title
            title = doc.add_heading('Analysis Report', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Process the analysis text and format properly
            sections = analysis_text.split("\n\n")
            
            for section in sections:
                if not section.strip():
                    continue
                    
                lines = section.split("\n")
                if not lines:
                    continue
                    
                # Check if this is a section header - improved detection for markdown
                first_line = lines[0].strip()
                
                if (first_line.startswith("#") or 
                    first_line.startswith("**") and first_line.endswith("**") or
                    first_line.isupper() or 
                    (len(first_line) < 100 and first_line.endswith(":")) or
                    re.match(r'^\d+\.\s*\*\*.*\*\*', first_line)):  # Numbered sections with bold
                    # This is likely a heading
                    heading_text = (first_line
                                  .replace("#", "")
                                  .replace("**", "")  # Remove markdown bold
                                  .strip()
                                  .rstrip(":"))
                    
                    # Remove numbering if present
                    heading_text = re.sub(r'^\d+\.\s*', '', heading_text)
                    
                    doc.add_heading(heading_text, level=1)
                    
                    # Add the rest of the lines as content
                    content_lines = lines[1:]
                else:
                    # This is regular content
                    content_lines = lines
                
                # Process content lines
                for line in content_lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    # Check if this is a bullet point
                    if line.startswith(("-", "*", "•")) or re.match(r'^\d+\.', line):
                        # Clean bullet text and add as bullet point
                        bullet_text = re.sub(r'^[-*•\d+\.]\s*', '', line)
                        # Remove markdown formatting for DOCX
                        bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Remove bold
                        bullet_text = re.sub(r'\*(.*?)\*', r'\1', bullet_text)      # Remove italic
                        p = doc.add_paragraph(bullet_text, style='List Bullet')
                    else:
                        # Add as regular paragraph, cleaning markdown
                        clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove bold markdown
                        clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Remove italic markdown
                        doc.add_paragraph(clean_line)
                
                # Add spacing between sections
                doc.add_paragraph()

            # Save DOCX to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_content = buffer.getvalue()

            # Upload DOCX
            container_client = self.blob_service_client.get_container_client(
                self.config.storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_url)

            blob_client.upload_blob(docx_content, overwrite=True)
            return blob_client.url

        except Exception as e:
            logger.error(f"Error generating/uploading DOCX: {str(e)}")
            raise StorageServiceError(f"Error generating/uploading DOCX: {str(e)}") from e
