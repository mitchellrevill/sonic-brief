import os
import logging
from typing import Optional, AsyncGenerator
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from azure.storage.blob.aio import BlobClient as AsyncBlobClient
from azure.identity import DefaultAzureCredential
from azure.core.exceptions import AzureError
from datetime import datetime, timedelta
from urllib.parse import urlparse
from azure.core.exceptions import ResourceNotFoundError
from ...core.config import AppConfig


class StorageService:
    def __init__(self, config: AppConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Prefer key-based authentication for local development
        # Falls back to managed identity for cloud deployment
        if config.azure_storage_key:
            self.logger.info("Using Azure Storage key-based authentication")
            self.credential = config.azure_storage_key
        else:
            self.logger.info("Using Azure Storage managed identity authentication")
            self.credential = DefaultAzureCredential()

        # Initialize blob service client
        self.blob_service_client = BlobServiceClient(
            account_url=self.config.azure_storage_account_url, credential=self.credential
        )

    def generate_sas_token(self, blob_url: str) -> Optional[str]:
        """Generate SAS token for a blob URL using managed identity"""
        try:
            if not blob_url:
                return None

            # Parse blob URL to get container and blob name
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.strip("/").split("/")
            if len(path_parts) < 2:
                return None

            container_name = path_parts[0]
            blob_name = "/".join(path_parts[1:])

            # Use different SAS generation method depending on authentication type
            if isinstance(self.credential, str):
                # Key-based authentication - use account key
                sas_token = generate_blob_sas(
                    account_name=parsed_url.netloc.split(".")[0],
                    container_name=container_name,
                    blob_name=blob_name,
                    account_key=self.credential,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=8),
                )
            else:
                # Managed identity - get user delegation key
                user_delegation_key = self.blob_service_client.get_user_delegation_key(
                    key_start_time=datetime.utcnow() - timedelta(minutes=5),
                    key_expiry_time=datetime.utcnow() + timedelta(hours=8),
                )

                # Generate SAS token using user delegation key
                sas_token = generate_blob_sas(
                    account_name=parsed_url.netloc.split(".")[0],
                    container_name=container_name,
                    blob_name=blob_name,
                    user_delegation_key=user_delegation_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=8),
                )

            return sas_token

        except Exception as e:
            self.logger.error(f"Error generating SAS token: {str(e)}")
            return None

    def add_sas_token_to_url(self, blob_url: str) -> str:
        """Add SAS token to blob URL if not already present"""
        if not blob_url:
            return blob_url

        sas_token = self.generate_sas_token(blob_url)
        if sas_token:
            self.logger.debug(f"Adding SAS token to blob URL: {blob_url}")
            return f"{blob_url}?{sas_token}"
        self.logger.debug(f"No SAS token generated for blob URL: {blob_url}")
        return blob_url

    def upload_file(self, file_path: str, original_filename: str) -> str:
        """Upload a file to blob storage"""
        try:
            container_client = self.blob_service_client.get_container_client(
                self.config.azure_storage_recordings_container
            )            # Sanitize filename - replace spaces with underscores
            sanitized_filename = original_filename.replace(" ", "_")
            self.logger.debug(
                f"Sanitized filename: {original_filename} -> {sanitized_filename}"
            )
            
            # Generate blob name with date and nested structure including timestamp for uniqueness
            current_date = datetime.now().strftime("%Y-%m-%d")
            timestamp = datetime.now().strftime("%H%M%S_%f")[:-3]  # HHMMSS_milliseconds
            file_name_without_ext = os.path.splitext(sanitized_filename)[0]
            # Include timestamp in both folder and filename to ensure uniqueness
            blob_name = f"{current_date}/{file_name_without_ext}_{timestamp}/{sanitized_filename}"

            blob_client = container_client.get_blob_client(blob_name)

            # Upload the file
            self.logger.info(f"Uploading file to blob storage: {blob_name}")
            with open(file_path, "rb") as data:
                blob_client.upload_blob(data, overwrite=True)

            return blob_client.url

        except AzureError as e:
            self.logger.error(f"Azure storage error: {str(e)}")
            raise
        except Exception as e:
            self.logger.error(f"Error uploading file: {str(e)}")
            raise

    def generate_and_upload_docx(self, analysis_text: str, blob_name: str) -> str:
        """Generate a DOCX from analysis text and upload to blob storage. Return the blob URL."""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import io
            import re

            # Create DOCX in memory
            doc = Document()
            
            # Add title
            title = doc.add_heading('Analysis Report', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Process the analysis text and format properly
            sections = analysis_text.split("\n\n")
            
            for section in sections:
                if not section.strip():
                    continue
                    
                lines = section.split("\n")
                if not lines:
                    continue
                    
                # Check if this is a section header - improved detection for markdown
                first_line = lines[0].strip()
                
                if (first_line.startswith("#") or 
                    first_line.startswith("**") and first_line.endswith("**") or
                    first_line.isupper() or 
                    (len(first_line) < 100 and first_line.endswith(":")) or
                    re.match(r'^\d+\.\s*\*\*.*\*\*', first_line)):  # Numbered sections with bold
                    # This is likely a heading
                    heading_text = (first_line
                                  .replace("#", "")
                                  .replace("**", "")  # Remove markdown bold
                                  .strip()
                                  .rstrip(":"))
                    
                    # Remove numbering if present
                    heading_text = re.sub(r'^\d+\.\s*', '', heading_text)
                    
                    doc.add_heading(heading_text, level=1)
                    
                    # Add the rest of the lines as content
                    content_lines = lines[1:]
                else:
                    # This is regular content
                    content_lines = lines
                
                # Process content lines
                for line in content_lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    # Check if this is a bullet point
                    if line.startswith(("-", "*", "•")) or re.match(r'^\d+\.', line):
                        # Clean bullet text and add as bullet point
                        bullet_text = re.sub(r'^[-*•\d+\.]\s*', '', line)
                        # Remove markdown formatting for DOCX
                        bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Remove bold
                        bullet_text = re.sub(r'\*(.*?)\*', r'\1', bullet_text)      # Remove italic
                        p = doc.add_paragraph(bullet_text, style='List Bullet')
                    else:
                        # Add as regular paragraph, cleaning markdown
                        clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove bold markdown
                        clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)  # Remove italic markdown
                        doc.add_paragraph(clean_line)
                
                # Add spacing between sections
                doc.add_paragraph()

            # Save DOCX to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_content = buffer.getvalue()

            # Upload DOCX
            container_client = self.blob_service_client.get_container_client(
                self.config.azure_storage_recordings_container
            )
            blob_client = container_client.get_blob_client(blob_name)

            self.logger.info(f"Uploading DOCX to blob storage: {blob_name}")
            blob_client.upload_blob(docx_content, overwrite=True)
            return blob_client.url

        except Exception as e:
            self.logger.error(f"Error generating/uploading DOCX: {str(e)}")
            raise

    async def stream_blob_content(
        self, file_blob_url: str
    ) -> AsyncGenerator[bytes, None]:
        """
        Stream content from a blob asynchronously.

        Args:
            file_blob_url (str): URL of the blob to stream.

        Returns:
            AsyncGenerator[bytes, None]: Yields chunks of file content asynchronously.

        Raises:
            ValueError: If the provided URL is invalid or missing required parts.
            ResourceNotFoundError: If the blob does not exist.
            Exception: For other unexpected errors.
        """
        if not file_blob_url:
            raise ValueError("Blob URL cannot be empty.")

        try:
            parsed_url = urlparse(file_blob_url)
            if not parsed_url.path:
                raise ValueError("Invalid blob URL: Missing path.")

            # Extract the blob name from the URL
            # First try the expected recordings container
            if self.config.azure_storage_recordings_container in parsed_url.path:
                blob_name = parsed_url.path.split(
                    self.config.azure_storage_recordings_container, 1
                )[-1].lstrip("/")
            else:
                # For transcription files or other assets that might be in different containers,
                # try to extract container and blob name from the URL path
                path_parts = parsed_url.path.strip('/').split('/')
                if len(path_parts) >= 2:
                    # Assume format: /container_name/blob_name or /container_name/folder/blob_name
                    container_name = path_parts[0]
                    blob_name = '/'.join(path_parts[1:])
                    self.logger.warning(f"Blob URL uses different container '{container_name}' instead of expected '{self.config.azure_storage_recordings_container}'. Using container: {container_name}")
                else:
                    raise ValueError(f"Blob URL path format not recognized: {parsed_url.path}")
            
            if not blob_name:
                raise ValueError("Could not extract blob name from URL")
            self.logger.debug(f"Extracted blob name: {blob_name}")

            # Create an async blob client
            async_blob_client = AsyncBlobClient(
                account_url=self.config.azure_storage_account_url,
                container_name=self.config.azure_storage_recordings_container,
                blob_name=blob_name,
                credential=self.credential,
            )

            # Stream the blob content in chunks
            async with async_blob_client:
                # Get the downloader without specifying chunk size
                # The chunks() method doesn't accept a chunk_size parameter
                downloader = await async_blob_client.download_blob()

                # Stream the chunks as they come
                async for chunk in downloader.chunks():
                    yield chunk

        except ValueError as ve:
            self.logger.warning(f"Validation error: {ve}")
            raise
        except ResourceNotFoundError as rnfe:
            self.logger.error(f"Blob not found: {rnfe}")
            raise
        except Exception as e:
            self.logger.error(
                f"Unexpected error streaming blob content: {str(e)}", exc_info=True
            )
            raise
